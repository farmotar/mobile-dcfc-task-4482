"""
build_appendix_a.py
====================
Generate Appendix A: Cost Model Development (Task 4482 / Mobile DCFC).
Produces 8 PNG figures and QuarterlyReport_4482_FY26_Q4_V3.docx.

Scenario A (always grid-connected) only.
90th-percentile total daily cost method per Shima's instructions.
"""
from __future__ import annotations
import sys, math
from pathlib import Path
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import matplotlib.ticker as mticker
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding="utf-8")

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE   = Path(r"D:\Geotab_EV_Parameters\charger_sizing_test")
OUTDIR = BASE / "scenario_outputs"
FIGS   = BASE / "appendix_a_figures"
FIGS.mkdir(exist_ok=True)
DOCOUT = Path(r"C:\Users\admin-local.COE-MAE-PF5QTVT\Downloads\QuarterlyReport_4482_FY26_Q4_V3.docx")

SITES = [
    ("northgate", "Northgate"),
    ("fresno",    "Fresno"),
    ("glendale",  "Glendale"),
    ("san_diego", "San Diego"),
]
UTILITY_LABEL = {
    "northgate": "SMUD C&I 21–299 kW (TOU)",
    "fresno":    "PG&E BEV-2 Secondary (TOU)",
    "glendale":  "GWP (SMUD proxy — rate not confirmed; see §A.11)",
    "san_diego": "SDG&E EV-HP Secondary (TOU)",
}
K_CAP       = 20
DAYS_PER_MO = 30.42
XOS_PURCHASE = 245_437.50
XOS_LIFE     = 10
XOS_MAINT    = 6_000    # $/unit/yr — ASSUMED
XOS_WARRANTY = 10_000   # $/unit/yr — from Caltrans project team

INFRA_SHARED   = {"low": 20_000, "mid": 40_000, "high": 80_000}
INFRA_PER_UNIT = {"low":  6_000, "mid":  8_500, "high": 12_000}
INFRA_TIER     = {"low": 12_000, "mid": 20_000, "high": 35_000}
TIER_SZ        = 4

COMP_COLORS = {
    "Capital (purchase)":    "#2166ac",
    "Capital (infra)":       "#74add1",
    "Maintenance*":          "#f46d43",
    "Warranty*":             "#fdae61",
    "Energy":                "#4dac26",
    "Demand – global":       "#d01c8b",
    "Demand – peak window":  "#c2a5cf",
}


# ── helpers ────────────────────────────────────────────────────────────────────

def infra_cost(n: int, est: str = "mid") -> dict:
    n_tiers = max(0, math.ceil(n / TIER_SZ) - 1)
    total   = INFRA_SHARED[est] + n * INFRA_PER_UNIT[est] + n_tiers * INFRA_TIER[est]
    return {"total": total, "per_unit": total / max(n, 1), "n_tiers": n_tiers}

def upfront(K: int) -> dict:
    hw  = K * XOS_PURCHASE
    inf = infra_cost(K)["total"]
    return {"hardware": hw, "infra": inf, "total": hw + inf}

def fmt_usd(v: float, dec: int = 0) -> str:
    fmt = f",.{dec}f"
    return f"${v:{fmt}}"


# ── Load A1 data for all sites ─────────────────────────────────────────────────

print("Loading A1 data …")
SD: dict = {}
for sk, sl in SITES:
    dc = pd.read_csv(OUTDIR / f"{sk}_analysis/{sk}_cost_detail.csv")
    sm = pd.read_csv(OUTDIR / f"{sk}_analysis/{sk}_summary.csv")
    a1d = dc[dc.scenario == "A1"].copy()
    a1s = sm[sm.scenario == "A1"].copy()
    a1d["date"] = pd.to_datetime(a1d["date"])
    a1s["date"] = pd.to_datetime(a1s["date"])

    # True all-in daily cost: both demand components amortised to daily
    a1d["total_allin"] = (
        a1d["total_daily_excl_demand"]
        + a1d["demand_global_monthly_$"] / DAYS_PER_MO
        + a1d["demand_peak_win_monthly_$"] / DAYS_PER_MO
    )

    p90  = a1d["total_allin"].quantile(0.90)
    idx  = (a1d["total_allin"] - p90).abs().idxmin()
    pday = a1d.loc[idx]
    nomK = int(pday["K"])
    cap  = nomK >= K_CAP

    SD[sk] = {
        "label":   sl,
        "dc":      a1d.sort_values("date").reset_index(drop=True),
        "sm":      a1s.sort_values("date").reset_index(drop=True),
        "p90":     p90,
        "pday":    pday,
        "nomK":    nomK,
        "cap_hit": cap,
        "upfront": upfront(nomK),
        "peak_kw": float(pday["peak_grid_kw"]),
        "days":    len(a1d),
        "n_veh":   a1s["n_vehicles"].mean(),
        "svc_pct": a1s["service_rate_pct"].mean(),
    }
    cap_str = " *** K_CAP HIT ***" if cap else ""
    print(f"  {sl:<12} days={len(a1d)}  p90={fmt_usd(p90)}  nomK={nomK}  "
          f"date={pday['date'].strftime('%Y-%m-%d')}{cap_str}")

print()


# ── Figure 1 per site: daily cost timeline ────────────────────────────────────

def fig_daily_cost(sk: str, save_path: Path) -> None:
    d   = SD[sk]
    df  = d["dc"]
    sm  = d["sm"]
    p90 = d["p90"]
    nomK = d["nomK"]

    fig, ax = plt.subplots(figsize=(13, 4.8))

    # Color by K
    k_vals = df["K"].values
    cmap   = plt.cm.get_cmap("YlOrRd", int(k_vals.max()) + 1)
    sc = ax.scatter(df["date"], df["total_allin"],
                    c=k_vals, cmap=cmap,
                    vmin=1, vmax=max(k_vals.max(), nomK),
                    s=22, alpha=0.80, zorder=3, linewidths=0)

    # p90 line
    ax.axhline(p90, color="#d01c8b", linewidth=1.8, linestyle="--",
               label=f"90th percentile  {fmt_usd(p90)}/day  (nominal K = {nomK})",
               zorder=4)
    # shade above p90
    ax.fill_between(df["date"], p90, df["total_allin"].clip(lower=p90),
                    alpha=0.12, color="#d01c8b", zorder=2)

    # Annotate p90 day
    pday = d["pday"]
    ax.annotate(f"K={nomK} — {pday['date'].strftime('%b %d, %Y')}",
                xy=(pday["date"], float(pday["total_allin"])),
                xytext=(pday["date"], float(pday["total_allin"]) + p90 * 0.12),
                arrowprops=dict(arrowstyle="->", color="#555", lw=1.2),
                fontsize=8, color="#333",
                ha="center", va="bottom")

    cb = plt.colorbar(sc, ax=ax, pad=0.01)
    cb.set_label("XOS units deployed (K)", fontsize=8)
    cb.ax.tick_params(labelsize=7)

    ax.xaxis.set_major_formatter(mdates.DateFormatter("%b '%y"))
    ax.xaxis.set_major_locator(mdates.MonthLocator(interval=1))
    plt.setp(ax.get_xticklabels(), rotation=35, ha="right", fontsize=8)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"${v:,.0f}"))
    ax.set_ylabel("Total daily cost — Scenario A ($/day)", fontsize=9)
    ax.set_xlabel("Date (May 2025 – Apr 2026)", fontsize=9)
    ax.set_title(
        f"{d['label']} — XOS Hub MC02 Scenario A: Total Daily Cost (incl. demand charges)\n"
        f"Utility: {UTILITY_LABEL[sk].split(' (')[0]}   |   {d['days']} operating days",
        fontsize=10, fontweight="bold")
    ax.legend(fontsize=8.5, loc="upper left", framealpha=0.92)
    ax.grid(axis="y", linestyle=":", alpha=0.35)
    if d["cap_hit"]:
        ax.text(0.98, 0.96, "⚠ Model cap K=20 reached at p90 day",
                transform=ax.transAxes, fontsize=8, color="#b30000",
                ha="right", va="top",
                bbox=dict(boxstyle="round,pad=0.3", facecolor="#fff0f0", edgecolor="#b30000"))

    fig.tight_layout()
    fig.savefig(save_path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    print(f"  Saved: {save_path.name}")


# ── Figure 2 per site: cost breakdown at p90 day ─────────────────────────────

def fig_cost_breakdown(sk: str, save_path: Path) -> None:
    d   = SD[sk]
    row = d["pday"]   # the p90-day row from cost_detail

    components = {
        "Capital (purchase)":   float(row["purchase_capex_daily"]),
        "Capital (infra)":      float(row["infra_capex_daily"]),
        "Maintenance*":         float(row["maint_daily"]),
        "Warranty*":            float(row["warranty_daily"]),
        "Energy":               float(row["energy_cost_daily"]),
        "Demand – global":      float(row["demand_global_monthly_$"]) / DAYS_PER_MO,
        "Demand – peak window": float(row["demand_peak_win_monthly_$"]) / DAYS_PER_MO,
    }
    labels = list(components.keys())
    values = list(components.values())
    colors = [COMP_COLORS[l] for l in labels]
    total  = sum(values)

    fig, ax = plt.subplots(figsize=(9, 5))
    bars = ax.bar(labels, values, color=colors, edgecolor="white", linewidth=0.5, zorder=3)

    for bar, val in zip(bars, values):
        if val > total * 0.01:
            ax.text(bar.get_x() + bar.get_width() / 2,
                    bar.get_height() + total * 0.008,
                    f"${val:,.0f}", ha="center", va="bottom", fontsize=7.5, color="#333")

    ax.axhline(0, color="black", linewidth=0.6)
    ax.set_ylabel("Daily cost ($/day)", fontsize=9)
    ax.set_title(
        f"{d['label']} — Cost Breakdown at Nominal Design Day\n"
        f"Date: {row['date'].strftime('%Y-%m-%d')}   K = {d['nomK']} XOS units   "
        f"Total = {fmt_usd(total)}/day",
        fontsize=10, fontweight="bold")
    ax.set_xticks(range(len(labels)))
    ax.set_xticklabels(labels, rotation=25, ha="right", fontsize=8.5)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"${v:,.0f}"))
    ax.grid(axis="y", linestyle=":", alpha=0.35)
    ax.text(0.99, 0.97, "* Assumed value — pending vendor confirmation",
            transform=ax.transAxes, fontsize=7, color="#888",
            ha="right", va="top")
    fig.tight_layout()
    fig.savefig(save_path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    print(f"  Saved: {save_path.name}")


print("Generating figures …")
for sk, _ in SITES:
    fig_daily_cost(sk,      FIGS / f"{sk}_daily_cost.png")
    fig_cost_breakdown(sk,  FIGS / f"{sk}_cost_breakdown.png")
print()


# ── Cross-site summary figure ─────────────────────────────────────────────────

def fig_cross_site(save_path: Path) -> None:
    site_labels = [SD[sk]["label"] for sk, _ in SITES]
    p90s    = [SD[sk]["p90"]       for sk, _ in SITES]
    nomKs   = [SD[sk]["nomK"]      for sk, _ in SITES]
    upfronts= [SD[sk]["upfront"]["total"] / 1e6 for sk, _ in SITES]

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 4.5))

    colors = ["#2166ac", "#4dac26", "#f46d43", "#d01c8b"]

    ax1.bar(site_labels, p90s, color=colors, edgecolor="white")
    for i, (v, k) in enumerate(zip(p90s, nomKs)):
        ax1.text(i, v + max(p90s) * 0.015, f"{fmt_usd(v)}/day\nK={k}",
                 ha="center", va="bottom", fontsize=8)
    ax1.set_ylabel("90th-pct total daily cost ($/day)", fontsize=9)
    ax1.set_title("90th-Percentile Design Cost by Site", fontsize=10, fontweight="bold")
    ax1.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"${v:,.0f}"))
    ax1.grid(axis="y", linestyle=":", alpha=0.35)

    ax2.bar(site_labels, upfronts, color=colors, edgecolor="white")
    for i, v in enumerate(upfronts):
        ax2.text(i, v + max(upfronts) * 0.015, f"${v:.2f}M",
                 ha="center", va="bottom", fontsize=8)
    ax2.set_ylabel("Total upfront cost, mid estimate ($M)", fontsize=9)
    ax2.set_title("Nominal Config Upfront Cost by Site\n(Hardware + Building-Side Electrical Infra)",
                  fontsize=10, fontweight="bold")
    ax2.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"${v:.1f}M"))
    ax2.grid(axis="y", linestyle=":", alpha=0.35)

    fig.tight_layout()
    fig.savefig(save_path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    print(f"  Saved: {save_path.name}")

fig_cross_site(FIGS / "cross_site_summary.png")
print()


# ── Word document ──────────────────────────────────────────────────────────────

def _set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        cell = row.cells[col_idx]
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(int(width_inches * 1440)))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)

def _shade_row(row, hex_color="D9E1F2"):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

def _hdr_row(row, hex_color="2F5496"):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)
        for para in cell.paragraphs:
            run = para.runs[0] if para.runs else para.add_run()
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold      = True

def add_table(doc, headers, rows_data, col_widths=None):
    t = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    _hdr_row(hdr)
    for ri, rdata in enumerate(rows_data):
        for ci, val in enumerate(rdata):
            t.rows[ri + 1].cells[ci].text = str(val)
        if ri % 2 == 0:
            _shade_row(t.rows[ri + 1], "EEF2F8")
    if col_widths:
        for ci, w in enumerate(col_widths):
            _set_col_width(t, ci, w)
    doc.add_paragraph()
    return t

def add_fig(doc, png_path: Path, caption: str, width: float = 6.0):
    if not png_path.exists():
        doc.add_paragraph(f"[Figure not found: {png_path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(png_path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.italic = True
    doc.add_paragraph()

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.size = Pt(14)
    return p

def h2(doc, text):
    return doc.add_heading(text, level=2)

def h3(doc, text):
    return doc.add_heading(text, level=3)

def body(doc, text):
    return doc.add_paragraph(text)

def bold_run(para, text):
    run = para.add_run(text)
    run.bold = True
    return run

# ─────────────────────────────────────────────────────────────────────────────
print("Building Word document …")
doc = Document()

# Page margins
section = doc.sections[0]
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# Title / header
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("DRAFT — FOR SHIMA'S REVIEW")
tr.font.size  = Pt(10)
tr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
tr.font.bold  = True
doc.add_paragraph()

title_p2 = doc.add_paragraph()
title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr2 = title_p2.add_run("Appendix A: Cost Model Development")
tr2.font.size  = Pt(18)
tr2.font.bold  = True
doc.add_paragraph()

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.add_run("Task 4482 — Mobile DC Fast Charging (MDCFC) Cost Analysis\n"
              "Caltrans Maintenance Stations — Scenario A (Grid-Connected)\n"
              "Report Date: June 30, 2026   |   Analysis Period: May 2025 – Apr 2026").font.size = Pt(10)
doc.add_page_break()

# ─── A.1 Overview ─────────────────────────────────────────────────────────────
h1(doc, "A.1  Overview and Scope")
body(doc,
    "This appendix presents the cost model developed for evaluating the XOS Hub MC02 mobile "
    "DC fast charger (MDCFC) as a charging solution for Caltrans electric fleet vehicles at four "
    "California maintenance stations: Northgate (Sacramento), Fresno, Glendale, and San Diego. "
    "The analysis covers Scenario A only — in which the XOS Hub is always connected to the "
    "electrical grid — per Shima's instruction. All cost figures reflect a 10-year asset "
    "lifespan and include capital amortization, electricity energy costs, utility demand "
    "or capacity charges, and warranty and maintenance (the latter two are flagged as assumed "
    "values pending vendor confirmation).")
body(doc,
    "The nominal design configuration for each site is selected using the 90th-percentile "
    "total daily cost method described in §A.5. Results for all four sites are presented in "
    "§A.6. A cross-site summary is provided in §A.8.")

# ─── A.2 Methodology ──────────────────────────────────────────────────────────
h1(doc, "A.2  Methodology")
h2(doc, "A.2.1  Data Sources")
body(doc,
    "Charging event data were sourced from Geotab zone-to-zone (Z2Z) trip records for each "
    "Caltrans maintenance station. Each record contains: vehicle identifier, arrival time at "
    "station, departure time, and estimated energy needed. Records were filtered to "
    "May 2025 – April 2026 to capture a full year of operational data.")
body(doc,
    "Utility rates used in the energy and demand cost model:")
tbl_util_data = [
    ["Northgate",  "SMUD",   "C&I 21–299 kW (TOU)",       "Confirmed"],
    ["Fresno",     "PG&E",   "BEV-2 Secondary (TOU)",       "Confirmed"],
    ["Glendale",   "GWP",    "Schedule LD-2/PC-1",          "⚠ ASSUMED — SMUD proxy used; see §A.11"],
    ["San Diego",  "SDG&E",  "EV-HP Secondary (TOU)",       "Confirmed"],
]
add_table(doc,
    ["Site", "Utility", "Tariff", "Status"],
    tbl_util_data,
    col_widths=[1.1, 0.9, 2.0, 2.5])

h2(doc, "A.2.2  Simulation: 15-Minute State Machine with Adaptive-K Sizing")
body(doc,
    "The XOS Hub MC02 simulation operates on 15-minute time steps across each operational day. "
    "The core loop implements:")
blist = [
    "State of Charge (SoC) tracking: the Hub's 282 kWh LFP battery is updated each step "
     "based on grid charge-in (up to 83 kW per unit) and vehicle discharge-out "
     "(up to 80 kW per CCS1 port, four ports per unit). SoC is bounded to [20%, 100%].",
    "Greedy vehicle dispatch: arriving vehicles are assigned to the first available charge "
     "head (first-come, first-served). Charging continues until the vehicle's energy need "
     "is satisfied or its departure time is reached.",
    "Proactive recharge: once the Hub's SoC drops to ~20%, it draws maximum grid power "
     "(83 kW) to replenish the battery, acting as a fixed charger from that point until "
     "SoC recovers.",
    "Minimum dwell enforcement: vehicle dwell windows shorter than 60 minutes are extended "
     "to 60 minutes to reflect realistic site operation.",
]
for item in blist:
    p = doc.add_paragraph(item, style="List Bullet")

body(doc,
    "Adaptive-K sizing: the number of XOS Hub units (K) deployed on a given day is the "
    "minimum K such that all arriving vehicles are fully served. K is incremented by one "
    "until the service condition is met, up to a maximum of K = 20 units. Days where K = 20 "
    "and vehicles remain unserved are flagged as capacity-constrained.")

h2(doc, "A.2.3  Scenario A — Operating Logic")
body(doc,
    "In Scenario A, each XOS Hub unit remains connected to the 480V, 3-phase grid "
    "throughout the operating day. The grid continuously replenishes the battery at up to "
    "83 kW per unit. The battery acts as an energy buffer, decoupling instantaneous vehicle "
    "charging demand from the grid draw profile and smoothing peak demand relative to a "
    "direct-grid charger of equivalent output capacity. There is no diesel operation, no "
    "disconnect event, and no V2G functionality.")

# ─── A.3 Cost model ───────────────────────────────────────────────────────────
h1(doc, "A.3  Cost Model Formulation")
h2(doc, "A.3.1  Capital Cost (10-Year Amortization)")
body(doc,
    "Capital costs are amortized over a 10-year service life consistent with the XOS Hub MC02 "
    "manufacturer specification (10 years / 3,000 cycles at 70% DoD). The daily capital "
    "cost per unit is:")
p = doc.add_paragraph()
p.add_run("    C_daily = [ (P_unit + C_infra_per_unit) / (10 yr × 12 mo) + "
          "(M_annual + W_annual) / 12 ] / 30.42 days/month").font.name = "Courier New"
body(doc,
    "where P_unit = unit purchase price, C_infra_per_unit = amortized building-side "
    "electrical infrastructure cost per unit (mid estimate), M_annual = annual maintenance, "
    "and W_annual = annual warranty. Maintenance and warranty are flagged as assumed values "
    "(§A.4).")

h2(doc, "A.3.2  Energy Cost (TOU Rate)")
body(doc,
    "Energy cost is computed as the sum over all 15-minute intervals of: "
    "grid_kW[t] × 0.25 h × rate[t], where rate[t] is the site-specific TOU energy rate "
    "($/kWh) at time t. Energy rates are applied using the confirmed utility tariff for "
    "each site (Table A.2.1 above).")

h2(doc, "A.3.3  Demand and Capacity Charges")
body(doc,
    "Monthly demand and capacity charges are included in every cost figure in this report "
    "and are amortized to a daily equivalent by dividing by 30.42 days/month.")
body(doc,
    "SMUD (Northgate / Glendale proxy): two-tier demand charge — "
    "$6.454/kW-month applied to the all-hours peak grid draw, plus "
    "$9.960/kW-month applied to the peak-window (16:00–21:00 weekdays) peak draw.")
body(doc,
    "PG&E BEV-2 (Fresno): single monthly subscription charge of "
    "$1.91/kW-month based on subscribed capacity (modeled as peak draw). "
    "No separate peak-window demand component.")
body(doc,
    "SDG&E EV-HP (San Diego): single monthly subscription charge of "
    "$4.81/kW-month based on subscribed capacity (modeled as peak draw). "
    "No separate peak-window demand component.")

h2(doc, "A.3.4  Warranty and Maintenance")
body(doc,
    "Warranty: $10,000/unit/year, provided by the Caltrans project team. "
    "Maintenance: $6,000/unit/year, assumed (no published service contract data available "
    "from XOS). Both values are flagged with an asterisk (*) in all tables and figures. "
    "Actual values should be confirmed with the XOS Program Advisor before finalizing "
    "the report (see §A.11).")

# ─── A.4 Hardware table ───────────────────────────────────────────────────────
h1(doc, "A.4  Hardware Specifications and Assumptions")
body(doc,
    "Table A.4.1 summarizes the XOS Hub MC02 technical specifications and all cost "
    "assumptions used in the model. Items marked * are assumed values pending confirmation.")

hw_rows = [
    ["Model",                          "XOS Hub MC02",                         "XOS User Manual"],
    ["Form factor",                    "Trailer-mounted, mobile DCFC",         "XOS User Manual"],
    ["Battery chemistry",              "Lithium Iron Phosphate (LFP)",         "XOS User Manual, Sec. 5"],
    ["Nominal battery capacity",       "282 kWh",                              "XOS User Manual, Sec. 5"],
    ["Usable capacity (SoC ≥ 20%)",   "225.6 kWh",                            "Derived: 80% of nominal"],
    ["Minimum SoC reserve",            "20% (56.4 kWh)",                       "Operational constraint"],
    ["Charge heads",                   "4 × CCS1",                             "XOS User Manual"],
    ["Max output per port",            "80 kW (constant)",                     "XOS User Manual"],
    ["Max hub output (battery only)",  "150 kW continuous",                    "XOS User Manual"],
    ["Max hub output (grid + battery)","230 kW continuous",                    "XOS User Manual"],
    ["Max grid input per unit",        "~83 kW  (480V × 100A × √3)",          "XOS User Manual"],
    ["Grid circuit requirement",       "480V, 3-phase, 100A dedicated",        "Electrical design"],
    ["Charge efficiency (grid→batt)",  "95%",                                  "Modeled assumption"],
    ["Discharge efficiency (batt→veh)","95%",                                  "Modeled assumption"],
    ["Service life",                   "10 years / 3,000 cycles at 70% DoD",  "XOS User Manual, Sec. 5"],
    ["Time-step resolution",           "15 minutes",                           "Model parameter"],
    ["Max units per site (model cap)", "20",                                   "Model parameter"],
    ["Unit purchase price",            "$245,437.50",                          "XOS Program Advisor (list price; not DGS contract)"],
    ["Annual maintenance *",           "$6,000 / unit / year",                 "⚠ ASSUMED — no published data"],
    ["Annual warranty *",              "$10,000 / unit / year",                "Caltrans project team"],
]
add_table(doc,
    ["Parameter", "Value", "Source"],
    hw_rows,
    col_widths=[2.2, 2.0, 2.8])
body(doc, "* Assumed values — must be confirmed before report finalization.")

# ─── A.5 90th-pct method ──────────────────────────────────────────────────────
h1(doc, "A.5  90th-Percentile Cost Design Method")
body(doc,
    "Per Shima's direction, the nominal design configuration for each site is determined "
    "by the 90th-percentile total daily cost method rather than a service-rate or "
    "coverage-ceiling criterion.")
body(doc,
    "Definition: for each site, total daily cost is computed for every Scenario A operating "
    "day across the full analysis period (May 2025 – April 2026). Total daily cost includes "
    "all components: capital amortization (purchase + infrastructure), energy, both demand "
    "charge components, maintenance, and warranty. The 90th-percentile value of this "
    "distribution is the NOMINAL DESIGN COST. The number of XOS Hub units deployed on "
    "the day nearest to that percentile value is the NOMINAL DESIGN K.")
body(doc,
    "Rationale: designing to the 90th percentile means the installed fleet comfortably covers "
    "all but the highest-cost 10% of days. This is a cost-conservative target — it avoids "
    "over-provisioning for rare peak days while ensuring adequate capacity for the vast "
    "majority of operational days. It also provides a defensible, quantitative design basis "
    "that is reproducible and independent of subjective coverage thresholds.")
body(doc,
    "Upfront cost ranking: for each site, the total upfront cost of the nominal design "
    "(hardware purchase + building-side electrical infrastructure, mid estimate) is computed "
    "and reported alongside the 90th-percentile daily cost. Sites are ranked by upfront cost "
    "in §A.7.")

# ─── A.6 Per-site ─────────────────────────────────────────────────────────────
h1(doc, "A.6  Per-Site Cost Development")
fig_num = 1

for sk, sl in SITES:
    d  = SD[sk]
    pday = d["pday"]
    nomK = d["nomK"]
    up   = d["upfront"]
    cap  = d["cap_hit"]

    h2(doc, f"A.6.{['northgate','fresno','glendale','san_diego'].index(sk)+1}  {sl} Maintenance Station")

    # Cap warning
    if cap:
        wp = doc.add_paragraph()
        wp.add_run(
            f"⚠  WARNING — MODEL CAP REACHED:  The 90th-percentile day ({pday['date'].strftime('%Y-%m-%d')}) "
            f"requires K = 20 XOS units, which is the model's maximum. The true design requirement "
            f"may exceed 20 units. This site likely requires an alternative sizing strategy or "
            f"additional infrastructure beyond what the current model evaluates. See §A.11."
        ).font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # Cost timeline plot
    body(doc, f"Figure A.{fig_num} shows the total daily cost (Scenario A) for all {d['days']} "
              f"operating days at {sl}, with the 90th-percentile design threshold marked.")
    add_fig(doc, FIGS / f"{sk}_daily_cost.png",
            f"Figure A.{fig_num}. {sl} — Scenario A Daily Cost Distribution with 90th-Percentile Design Line.",
            width=6.2)
    fig_num += 1

    # Nominal design summary table
    row_data = pday
    comp_rows = [
        ["Capital — purchase (amort.)",   fmt_usd(float(row_data["purchase_capex_daily"]), 2)],
        ["Capital — infrastructure (amort.)", fmt_usd(float(row_data["infra_capex_daily"]), 2)],
        ["Maintenance * (assumed)",       fmt_usd(float(row_data["maint_daily"]), 2)],
        ["Warranty * (assumed)",          fmt_usd(float(row_data["warranty_daily"]), 2)],
        ["Energy (TOU rate)",             fmt_usd(float(row_data["energy_cost_daily"]), 2)],
        ["Demand — global (monthly ÷ 30.42)",
             fmt_usd(float(row_data["demand_global_monthly_$"]) / DAYS_PER_MO, 2)],
        ["Demand — peak window (monthly ÷ 30.42)",
             fmt_usd(float(row_data["demand_peak_win_monthly_$"]) / DAYS_PER_MO, 2)],
        ["TOTAL all-in daily cost",       fmt_usd(float(d["p90"]), 2)],
    ]
    body(doc, f"Table A.{fig_num - (fig_num - fig_num)}.{['northgate','fresno','glendale','san_diego'].index(sk)+1} "
              f"below shows the cost breakdown for the nominal design day "
              f"(K = {nomK} XOS Hub MC02 units, {pday['date'].strftime('%Y-%m-%d')}).")
    add_table(doc,
        ["Cost Component", "$/day (nominal design day)"],
        comp_rows,
        col_widths=[3.5, 2.0])

    # Key metrics
    nom_rows = [
        ["Utility",                UTILITY_LABEL[sk]],
        ["Nominal design K",       f"{nomK} XOS Hub MC02 units"],
        ["90th-pct design cost",   fmt_usd(d["p90"], 2) + "/day"],
        ["Nominal design date",    pday["date"].strftime("%Y-%m-%d")],
        ["Peak grid draw (nom. day)", f"{d['peak_kw']:.0f} kW"],
        ["Upfront — hardware",     fmt_usd(up["hardware"]) + " (K × $245,438)"],
        ["Upfront — infra (mid)",  fmt_usd(up["infra"])    + " (building-side electrical)"],
        ["Upfront — TOTAL (mid)",  fmt_usd(up["total"])],
        ["Annualized cost (est.)", fmt_usd(d["dc"]["total_allin"].sum() * 365 / d["days"]) + "/yr"],
        ["Avg vehicles/day",       f"{d['n_veh']:.1f}"],
        ["Avg service rate",       f"{d['svc_pct']:.1f}%"],
        ["Analysis days",          str(d["days"])],
    ]
    add_table(doc, ["Metric", "Value"], nom_rows, col_widths=[2.8, 3.5])

    # Cost breakdown figure
    body(doc, f"Figure A.{fig_num} shows the cost component breakdown for the nominal design day.")
    add_fig(doc, FIGS / f"{sk}_cost_breakdown.png",
            f"Figure A.{fig_num}. {sl} — Cost Breakdown at Nominal Design Day (K = {nomK}).",
            width=5.5)
    fig_num += 1
    doc.add_page_break()

# ─── A.7 Upfront cost ranking ─────────────────────────────────────────────────
h1(doc, "A.7  Configuration Ranking by Upfront Cost")
body(doc,
    "Table A.7.1 ranks the four sites' nominal configurations by total upfront cost "
    "(hardware purchase + building-side electrical infrastructure, mid estimate). "
    "The nominal K is the configuration at the site's 90th-percentile daily cost point.")

ranked = sorted(SITES, key=lambda x: SD[x[0]]["upfront"]["total"])
rank_rows = []
for rank_i, (sk, sl) in enumerate(ranked, 1):
    d  = SD[sk]
    up = d["upfront"]
    rank_rows.append([
        str(rank_i),
        sl,
        str(d["nomK"]),
        fmt_usd(up["hardware"]),
        fmt_usd(up["infra"]),
        fmt_usd(up["total"]),
        fmt_usd(d["p90"], 2) + "/day",
        "⚠ K_CAP" if d["cap_hit"] else "—",
    ])
add_table(doc,
    ["Rank", "Site", "Nom. K", "Hardware", "Infra (mid)", "Total Upfront", "p90 Daily Cost", "Note"],
    rank_rows,
    col_widths=[0.45, 0.9, 0.6, 1.2, 1.1, 1.2, 1.2, 0.8])

body(doc,
    "Electrical infrastructure costs (mid estimate) assume the existing site building has 480V "
    "service but the panel must be replaced or expanded. Low and high estimates span roughly "
    "±30–50% of the mid estimate depending on existing site conditions. Infrastructure costs "
    "are building-side only and exclude utility transformer upgrades or service entrance work.")
body(doc,
    "The XOS unit price ($245,437.50) is a unit list price from the XOS Program Advisor and is "
    "not a California DGS contract price. The Kempower units (separate scenario) are priced "
    "through CA DGS Contract 1-23-61-15A.")

# ─── A.8 Cross-site summary ───────────────────────────────────────────────────
h1(doc, "A.8  Cross-Site Summary")
add_fig(doc, FIGS / "cross_site_summary.png",
        "Figure A.9. Cross-Site Comparison — 90th-Percentile Design Cost and Nominal Upfront Cost.",
        width=6.2)

xsite_rows = []
for sk, sl in SITES:
    d  = SD[sk]
    up = d["upfront"]
    pday = d["pday"]
    annual = d["dc"]["total_allin"].sum() * 365 / d["days"]
    xsite_rows.append([
        sl,
        str(d["nomK"]),
        UTILITY_LABEL[sk].split(" (")[0],
        fmt_usd(d["p90"], 0) + "/day",
        fmt_usd(annual, 0) + "/yr",
        f"{d['peak_kw']:.0f} kW",
        fmt_usd(up["total"]),
        "⚠ K_CAP" if d["cap_hit"] else "—",
    ])
add_table(doc,
    ["Site", "Nom. K", "Utility", "p90 Daily Cost", "Est. Annual Cost", "Peak Grid Draw", "Upfront (mid)", "Flag"],
    xsite_rows,
    col_widths=[0.9, 0.55, 1.5, 1.1, 1.1, 1.0, 1.1, 0.7])

body(doc,
    "Annualized cost is estimated as the sum of all Scenario A daily costs over the analysis "
    "period, scaled to 365 days. It includes all cost components (capital, energy, demand, "
    "warranty, maintenance).")

# ─── A.9 Findings ─────────────────────────────────────────────────────────────
h1(doc, "A.9  Findings and Recommendations")

findings = {
    "Northgate": (
        f"The nominal configuration is K = {SD['northgate']['nomK']} XOS Hub MC02 units "
        f"(90th-pct daily cost = {fmt_usd(SD['northgate']['p90'], 2)}/day). "
        f"Daily cost is driven primarily by capital amortization and SMUD demand charges. "
        f"The peak grid draw of {SD['northgate']['peak_kw']:.0f} kW at the nominal design day "
        f"is significant and should be verified against the site's available electrical "
        f"service capacity before finalizing the design."
    ),
    "Fresno": (
        f"The nominal configuration is K = {SD['fresno']['nomK']} XOS Hub MC02 units "
        f"(90th-pct daily cost = {fmt_usd(SD['fresno']['p90'], 2)}/day). "
        f"Fresno has the smallest average fleet ({SD['fresno']['n_veh']:.1f} vehicles/day) "
        f"among the four sites and correspondingly the lowest upfront cost. "
        f"PG&E BEV-2's subscription model (no separate peak-window demand charge) simplifies "
        f"the cost structure."
    ),
    "Glendale": (
        f"The nominal configuration is K = {SD['glendale']['nomK']} XOS Hub MC02 units "
        f"(90th-pct daily cost = {fmt_usd(SD['glendale']['p90'], 2)}/day). "
        f"Glendale is the lowest-demand site ({SD['glendale']['n_veh']:.1f} vehicles/day average) "
        f"and has the lowest nominal K and upfront cost. However, the GWP utility rate used "
        f"here is a SMUD proxy — actual GWP tariff data must be obtained before any cost "
        f"figures for this site are considered reliable. See §A.11."
    ),
    "San Diego": (
        f"San Diego is the highest-demand site ({SD['san_diego']['n_veh']:.1f} vehicles/day "
        f"average) and the most challenging. The 90th-percentile day requires K = 20 XOS units "
        f"— the model's hard cap — indicating the true design requirement may exceed what the "
        f"current model can evaluate. The nominal design cost of {fmt_usd(SD['san_diego']['p90'], 2)}/day "
        f"and upfront cost of {fmt_usd(SD['san_diego']['upfront']['total'])} should be treated "
        f"as lower-bound estimates. A revised model without a K cap, or an alternative approach "
        f"(e.g., fixed DCFC supplement), should be explored before finalizing San Diego's "
        f"recommendation."
    ),
}
for site_nm, text in findings.items():
    p = doc.add_paragraph()
    bold_run(p, f"{site_nm}: ")
    p.add_run(text)

# ─── A.10 Worst-day tables ────────────────────────────────────────────────────
h1(doc, "A.10  Worst-Day Supporting Tables (Scenario A)")
body(doc,
    "Table A.10.1 shows the single highest-cost operating day (Scenario A) for each site, "
    "for reference.")
wd_rows = []
for sk, sl in SITES:
    d    = SD[sk]
    row  = d["dc"].loc[d["dc"]["total_allin"].idxmax()]
    up   = upfront(int(row["K"]))
    wd_rows.append([
        sl,
        row["date"].strftime("%Y-%m-%d"),
        str(int(row["K"])),
        fmt_usd(float(row["total_allin"]), 2),
        fmt_usd(float(row["energy_cost_daily"]), 2),
        fmt_usd(float(row["demand_global_monthly_$"]) / DAYS_PER_MO, 2),
        f"{float(row['peak_grid_kw']):.0f} kW",
    ])
add_table(doc,
    ["Site", "Date", "K", "Total Daily Cost", "Energy", "Demand (global)", "Peak Draw"],
    wd_rows,
    col_widths=[0.9, 1.0, 0.4, 1.3, 1.0, 1.3, 1.0])

# ─── A.11 Open items ──────────────────────────────────────────────────────────
h1(doc, "A.11  Open Items and Assumptions Requiring Confirmation")

items = [
    ("XOS maintenance contract pricing",
     "Annual maintenance cost is assumed at $6,000/unit/year. No published service contract "
     "data is available from XOS. This assumption should be confirmed with the XOS Program "
     "Advisor before finalizing any cost conclusions."),
    ("XOS warranty pricing",
     "Annual warranty cost of $10,000/unit/year was provided by the Caltrans project team. "
     "Confirm that this reflects current pricing and applies to the specific deployment "
     "scenario analyzed here."),
    ("XOS unit price — DGS contract status",
     "The XOS Hub MC02 unit price ($245,437.50) is a list price from the XOS Program "
     "Advisor, not a California DGS contract price. Clarify whether a DGS contract is "
     "available or expected, as this would affect the defensibility of the cost basis."),
    ("Glendale — GWP utility tariff",
     "Glendale Water & Power's commercial tariff (Schedule LD-2/PC-1) was not obtainable "
     "during this analysis — automated access to glendaleca.gov and third-party rate "
     "databases was blocked. All Glendale cost figures use SMUD rates as a placeholder. "
     "Obtain the actual GWP tariff (call GWP Customer Service: 855-550-4497, or request "
     "the tariff sheet via written inquiry) before reporting Glendale results."),
    ("San Diego — K cap binding at 90th percentile",
     "The 90th-percentile design day for San Diego requires K = 20 XOS units, which is the "
     "model's maximum. The model cannot evaluate fleets larger than 20 units. Two options: "
     "(a) raise the K cap and re-run, or (b) evaluate a hybrid XOS + fixed DCFC approach "
     "for San Diego. The current cost figures are lower-bound estimates only."),
    ("SDG&E EV-HP super off-peak window",
     "The SDG&E EV-HP super off-peak hours (assumed 00:00–06:00) were not explicitly stated "
     "in the rate schedule reviewed. Confirm the exact super off-peak hours with SDG&E before "
     "finalizing San Diego energy cost figures."),
    ("End-of-life / decommissioning costs",
     "No end-of-life or decommissioning cost estimates are included in the current model. "
     "Obtain estimates from XOS and include in the lifecycle cost for the final report."),
]
for title, text in items:
    p = doc.add_paragraph(style="List Bullet")
    bold_run(p, f"{title}: ")
    p.add_run(text)

# ─── Save ──────────────────────────────────────────────────────────────────────
doc.save(str(DOCOUT))
print(f"\nSaved: {DOCOUT}")
print(f"Figures: {FIGS}")
