"""
build_appendix_a_v2.py
=======================
Rebuilt Appendix A for Task 4482 — table-first layout, Part I (XOS) + Part II (Kempower).
Saves QuarterlyReport_4482_FY26_Q4_V3.docx (overwrites previous draft).
"""
from __future__ import annotations
import sys, math
from pathlib import Path
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import matplotlib.ticker as mticker
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding="utf-8")

BASE   = Path(r"D:\Geotab_EV_Parameters\charger_sizing_test")
OUTDIR = BASE / "scenario_outputs"
FIGS   = BASE / "appendix_a_figures"
FIGS.mkdir(exist_ok=True)
DOCOUT = Path(r"C:\Users\admin-local.COE-MAE-PF5QTVT\Downloads\QuarterlyReport_4482_FY26_Q4_V3e.docx")

SITES = [("northgate","Northgate"),("fresno","Fresno"),("glendale","Glendale"),("san_diego","San Diego")]
SITE_IDX = {s:i for i,(s,_) in enumerate(SITES)}

UTILITY_SHORT = {"northgate":"SMUD","fresno":"PG&E BEV-2","glendale":"PG&E BEV-2 (GWP proxy)","san_diego":"SDG&E EV-HP"}
UTILITY_FULL  = {
    "northgate": "SMUD C&I 21–299 kW (TOU)",
    "fresno":    "PG&E BEV-2 Secondary (TOU)",
    "glendale":  "PG&E BEV-2 Secondary (proxy — actual GWP tariff not obtained; see §A.14)",
    "san_diego": "SDG&E EV-HP Secondary (TOU)",
}
KMP_SITE_STATUS = {"northgate":"Complete","fresno":"Complete — PG&E BEV-2 rates",
                   "glendale":"Complete — PG&E BEV-2 proxy (GWP tariff unconfirmed, see §A.14)","san_diego":"Not computed (paused)"}

DAYS_PER_MO = 30.42
K_CAP       = 20

# XOS unit economics
XOS_PURCHASE = 245_437.50; XOS_LIFE = 10; XOS_MAINT = 6_000; XOS_WARRANTY = 10_000
# Kempower DGS
KMP = {
    "50kW":  {"purchase":23_408,"install":855,  "maint":1_573,"warranty":2_000,"life":8,"power":50},
    "150kW": {"purchase":62_154,"install":4_750,"maint":1_573,"warranty":2_000,"life":8,"power":150},
    "250kW": {"purchase":101_946,"install":5_225,"maint":1_573,"warranty":2_000,"life":8,"power":250},
}

INFRA_S = {"low":20_000,"mid":40_000,"high":80_000}
INFRA_U = {"low":6_000, "mid":8_500, "high":12_000}
INFRA_T = {"low":12_000,"mid":20_000,"high":35_000}
TIER_SZ = 4

def infra(n,e="mid"):
    nt=max(0,math.ceil(n/TIER_SZ)-1)
    tot=INFRA_S[e]+n*INFRA_U[e]+nt*INFRA_T[e]
    return {"total":tot,"per_unit":tot/max(n,1)}

def xos_upfront(K):
    hw=K*XOS_PURCHASE; inf=infra(K)["total"]
    return {"hw":hw,"inf":inf,"total":hw+inf}

def kmp_daily(t):  # per charger per day
    mc=(t["purchase"]+t["install"])/(t["life"]*12)
    mm=(t["maint"]+t["warranty"])/12
    return (mc+mm)/DAYS_PER_MO

KMP_DAILY = {k:kmp_daily(v) for k,v in KMP.items()}
fmt = lambda v,d=0: f"${v:,.{d}f}"


# ─────────────────────────────────────────────────────────────────────
# LOAD XOS A1 DATA
# ─────────────────────────────────────────────────────────────────────
print("Loading XOS A1 data …")
XD = {}
for sk,sl in SITES:
    dc=pd.read_csv(OUTDIR/f"{sk}_analysis/{sk}_cost_detail.csv")
    sm=pd.read_csv(OUTDIR/f"{sk}_analysis/{sk}_summary.csv")
    a1d=dc[dc.scenario=="A1"].copy(); a1s=sm[sm.scenario=="A1"].copy()
    a1d["date"]=pd.to_datetime(a1d["date"]); a1s["date"]=pd.to_datetime(a1s["date"])
    a1d["total_allin"]=(a1d["total_daily_excl_demand"]
        +a1d["demand_global_monthly_$"]/DAYS_PER_MO
        +a1d["demand_peak_win_monthly_$"]/DAYS_PER_MO)
    p90=a1d["total_allin"].quantile(0.90)
    idx=(a1d["total_allin"]-p90).abs().idxmin()
    pday=a1d.loc[idx]; nomK=int(pday["K"]); cap=nomK>=K_CAP
    XD[sk]={"label":sl,"dc":a1d.sort_values("date").reset_index(drop=True),
            "sm":a1s,"p90":p90,"pday":pday,"nomK":nomK,"cap_hit":cap,
            "upfront":xos_upfront(nomK),"peak_kw":float(pday["peak_grid_kw"]),
            "days":len(a1d),"n_veh":a1s.n_vehicles.mean(),"svc_pct":a1s.service_rate_pct.mean()}
    cap_s=" *** K_CAP ***" if cap else ""
    print(f"  XOS {sl:<12} days={len(a1d)}  p90={fmt(p90)}  nomK={nomK}{cap_s}")

# ─────────────────────────────────────────────────────────────────────
# LOAD KEMPOWER DATA (Northgate + Fresno complete; others partial/none)
# ─────────────────────────────────────────────────────────────────────
print("\nLoading Kempower data …")

def load_kmp_site(sk, per_day_subdir="kempower", summary_suffix="", label_suffix=""):
    summ_path = OUTDIR/f"{sk}_analysis/{sk}_kempower_summary{summary_suffix}.csv"
    if not summ_path.exists():
        print(f"  Kempower {sk}{label_suffix}: no summary file"); return None
    summ = pd.read_csv(summ_path)
    summ["date"]=pd.to_datetime(summ["date"])

    cb_files = sorted((OUTDIR/f"{sk}_analysis/per_day").glob(f"*/{per_day_subdir}/exact_milp_cost_breakdown.csv"))
    dem_rows = []
    for f in cb_files:
        date_str = f.parent.parent.name
        cb = pd.read_csv(f)
        row = {"date":pd.to_datetime(date_str)}
        for comp in ["global_demand_cost","peak_window_demand_cost","P_max_kw","P_peak_window_kw"]:
            r = cb[cb.component==comp]
            row[comp] = float(r["value"].iloc[0]) if not r.empty else 0.0
        dem_rows.append(row)
    dem_df = pd.DataFrame(dem_rows) if dem_rows else pd.DataFrame(columns=["date","global_demand_cost","peak_window_demand_cost"])

    if not dem_df.empty:
        summ = summ.merge(dem_df, on="date", how="left")
        summ["global_demand_cost"]      = summ.get("global_demand_cost",     pd.Series(0.0, index=summ.index)).fillna(0.0)
        summ["peak_window_demand_cost"] = summ.get("peak_window_demand_cost",pd.Series(0.0, index=summ.index)).fillna(0.0)
    else:
        summ["global_demand_cost"]      = 0.0
        summ["peak_window_demand_cost"] = 0.0

    summ["total_allin"] = (summ["capex_daily"] + summ["energy_cost"]
        + summ["global_demand_cost"]/DAYS_PER_MO
        + summ["peak_window_demand_cost"]/DAYS_PER_MO)

    p90 = summ["total_allin"].quantile(0.90)
    idx  = (summ["total_allin"]-p90).abs().idxmin()
    pday = summ.loc[idx]
    print(f"  Kempower {sk}{label_suffix:<12} days={len(summ)}  p90={fmt(p90)}  mix={pday['mix']}")
    return {"summ":summ,"p90":p90,"pday":pday}

KD = {}
for sk,_ in SITES:
    KD[sk] = load_kmp_site(sk)

# ─────────────────────────────────────────────────────────────────────
# LOAD GLENDALE SMUD SENSITIVITY DATA
# ─────────────────────────────────────────────────────────────────────
print("\nLoading Glendale SMUD sensitivity data …")

def load_xos_smud_glendale():
    path = OUTDIR/"glendale_analysis/glendale_cost_detail_smud.csv"
    if not path.exists():
        print("  XOS Glendale SMUD: no file found"); return None
    df = pd.read_csv(path)
    df["date"] = pd.to_datetime(df["date"])
    df["total_allin"] = (df["total_daily_excl_demand"]
        + df["demand_global_monthly_$"]/DAYS_PER_MO
        + df["demand_peak_win_monthly_$"]/DAYS_PER_MO)
    p90 = float(df["total_allin"].quantile(0.90))
    idx = (df["total_allin"]-p90).abs().idxmin()
    pday = df.loc[idx]
    print(f"  XOS Glendale (SMUD)  p90={fmt(p90)}  nomK={int(pday['K'])}")
    return {"dc":df.sort_values("date").reset_index(drop=True),"p90":p90,"pday":pday,"nomK":int(pday["K"])}

XD_GL_SMUD = load_xos_smud_glendale()
KD_GL_SMUD = load_kmp_site("glendale", per_day_subdir="kempower_smud",
                            summary_suffix="_smud", label_suffix=" (SMUD)")

# ─────────────────────────────────────────────────────────────────────
# FIGURES
# ─────────────────────────────────────────────────────────────────────
print("\nGenerating figures …")

SITE_COLORS = {"northgate":"#2166ac","fresno":"#4dac26","glendale":"#f46d43","san_diego":"#d01c8b"}

def fig_xos_daily(sk, path):
    d=XD[sk]; df=d["dc"]
    fig,ax=plt.subplots(figsize=(13,4.5))
    k_vals=df["K"].values
    cmap=matplotlib.colormaps.get_cmap("YlOrRd")
    sc=ax.scatter(df["date"],df["total_allin"],c=k_vals,cmap=cmap,
                  vmin=1,vmax=max(k_vals.max(),d["nomK"]),s=20,alpha=0.80,zorder=3,linewidths=0)
    ax.axhline(d["p90"],color="#d01c8b",linewidth=1.8,linestyle="--",
               label=f"90th percentile  {fmt(d['p90'])}/day  (K = {d['nomK']} units)",zorder=4)
    ax.fill_between(df["date"],d["p90"],df["total_allin"].clip(lower=d["p90"]),
                    alpha=0.10,color="#d01c8b",zorder=2)
    cb=plt.colorbar(sc,ax=ax,pad=0.01); cb.set_label("XOS units (K)",fontsize=8); cb.ax.tick_params(labelsize=7)
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%b '%y"))
    ax.xaxis.set_major_locator(mdates.MonthLocator(interval=1))
    plt.setp(ax.get_xticklabels(),rotation=35,ha="right",fontsize=8)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_:f"${v:,.0f}"))
    ax.set_ylabel("Total daily cost ($/day)",fontsize=9)
    ax.set_title(f"XOS Hub MC02 — {d['label']}: Daily Cost  ({UTILITY_SHORT[sk]})",
                 fontsize=10,fontweight="bold")
    ax.legend(fontsize=8.5,loc="upper left",framealpha=0.92)
    ax.grid(axis="y",linestyle=":",alpha=0.35)
    if d["cap_hit"]:
        ax.text(0.98,0.96,"⚠ Model cap K=20 at p90 day",transform=ax.transAxes,
                fontsize=8,color="#b30000",ha="right",va="top",
                bbox=dict(boxstyle="round,pad=0.3",facecolor="#fff0f0",edgecolor="#b30000"))
    fig.tight_layout(); fig.savefig(path,dpi=150,bbox_inches="tight"); plt.close(fig)
    print(f"  {path.name}")

def fig_xos_breakdown(sk, path):
    d=XD[sk]; row=d["pday"]
    comps={"Capital\n(purchase)":float(row["purchase_capex_daily"]),
           "Capital\n(infra)":float(row["infra_capex_daily"]),
           "Maint.*":float(row["maint_daily"]),
           "Warranty*":float(row["warranty_daily"]),
           "Energy":float(row["energy_cost_daily"]),
           "Demand\nglobal":float(row["demand_global_monthly_$"])/DAYS_PER_MO,
           "Demand\npeak-win":float(row["demand_peak_win_monthly_$"])/DAYS_PER_MO}
    colors=["#2166ac","#74add1","#f46d43","#fdae61","#4dac26","#d01c8b","#c2a5cf"]
    labels=list(comps.keys()); values=list(comps.values()); total=sum(values)
    fig,ax=plt.subplots(figsize=(9,4.5))
    bars=ax.bar(labels,values,color=colors,edgecolor="white",linewidth=0.5,zorder=3)
    for bar,val in zip(bars,values):
        if val>total*0.008:
            ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+total*0.008,
                    f"${val:,.0f}",ha="center",va="bottom",fontsize=7.5,color="#333")
    ax.set_title(f"XOS Hub MC02 — {d['label']}: Cost Breakdown at Nominal Design Day\n"
                 f"{row['date'].strftime('%Y-%m-%d')}   K={d['nomK']} units   Total={fmt(float(d['p90']))}/day",
                 fontsize=10,fontweight="bold")
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_:f"${v:,.0f}"))
    ax.set_xticks(range(len(labels))); ax.set_xticklabels(labels,fontsize=8.5)
    ax.grid(axis="y",linestyle=":",alpha=0.35)
    ax.text(0.99,0.97,"* Assumed value",transform=ax.transAxes,fontsize=7,color="#888",ha="right",va="top")
    fig.tight_layout(); fig.savefig(path,dpi=150,bbox_inches="tight"); plt.close(fig)
    print(f"  {path.name}")

def fig_kmp_daily(sk, path):
    kd=KD[sk]
    if kd is None: return
    df=kd["summ"].sort_values("date")
    fig,ax=plt.subplots(figsize=(13,4.5))
    ax.scatter(df["date"],df["total_allin"],color=SITE_COLORS[sk],s=20,alpha=0.75,zorder=3,linewidths=0)
    ax.axhline(kd["p90"],color="#555",linewidth=1.8,linestyle="--",
               label=f"90th percentile  {fmt(kd['p90'])}/day",zorder=4)
    ax.fill_between(df["date"],kd["p90"],df["total_allin"].clip(lower=kd["p90"]),
                    alpha=0.10,color="#555",zorder=2)
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%b '%y"))
    ax.xaxis.set_major_locator(mdates.MonthLocator(interval=1))
    plt.setp(ax.get_xticklabels(),rotation=35,ha="right",fontsize=8)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_:f"${v:,.0f}"))
    ax.set_ylabel("Total daily cost — Kempower ($/day)",fontsize=9)
    ax.set_title(f"Kempower Fixed DCFC — {SITES[SITE_IDX[sk]][1]}: Daily Cost  ({UTILITY_SHORT[sk]})",
                 fontsize=10,fontweight="bold")
    ax.legend(fontsize=8.5,loc="upper left",framealpha=0.92)
    ax.grid(axis="y",linestyle=":",alpha=0.35)
    fig.tight_layout(); fig.savefig(path,dpi=150,bbox_inches="tight"); plt.close(fig)
    print(f"  {path.name}")

def fig_kmp_breakdown(sk, path):
    kd=KD[sk]
    if kd is None: return
    row=kd["pday"]
    comps={"CapEx\n(amort.)":float(row["capex_daily"]),
           "Energy":float(row["energy_cost"]),
           "Demand\nglobal":float(row.get("global_demand_cost",0))/DAYS_PER_MO,
           "Demand\npeak-win":float(row.get("peak_window_demand_cost",0))/DAYS_PER_MO}
    colors=["#2166ac","#4dac26","#d01c8b","#c2a5cf"]
    labels=list(comps.keys()); values=list(comps.values()); total=sum(values)
    fig,ax=plt.subplots(figsize=(8,4.5))
    bars=ax.bar(labels,values,color=colors,edgecolor="white",linewidth=0.5,zorder=3)
    for bar,val in zip(bars,values):
        if val>total*0.008:
            ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+total*0.012,
                    f"${val:,.0f}",ha="center",va="bottom",fontsize=8,color="#333")
    sl=SITES[SITE_IDX[sk]][1]
    ax.set_title(f"Kempower — {sl}: Cost Breakdown at Nominal Design Day\n"
                 f"{row['date'].strftime('%Y-%m-%d')}   Mix: {row['mix']}   Total≈{fmt(total)}/day",
                 fontsize=10,fontweight="bold")
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_:f"${v:,.0f}"))
    ax.set_xticks(range(len(labels))); ax.set_xticklabels(labels,fontsize=9)
    ax.grid(axis="y",linestyle=":",alpha=0.35)
    fig.tight_layout(); fig.savefig(path,dpi=150,bbox_inches="tight"); plt.close(fig)
    print(f"  {path.name}")

def fig_xsite_xos(path):
    sls=[XD[sk]["label"] for sk,_ in SITES]
    p90s=[XD[sk]["p90"] for sk,_ in SITES]
    nomKs=[XD[sk]["nomK"] for sk,_ in SITES]
    ufs=[XD[sk]["upfront"]["total"]/1e6 for sk,_ in SITES]
    cols=[SITE_COLORS[sk] for sk,_ in SITES]
    fig,(ax1,ax2)=plt.subplots(1,2,figsize=(12,4.5))
    ax1.bar(sls,p90s,color=cols,edgecolor="white")
    for i,(v,k) in enumerate(zip(p90s,nomKs)):
        ax1.text(i,v+max(p90s)*0.015,f"{fmt(v)}\nK={k}",ha="center",va="bottom",fontsize=8)
    ax1.set_ylabel("90th-pct daily cost ($/day)",fontsize=9)
    ax1.set_title("XOS Hub — p90 Design Cost by Site",fontsize=10,fontweight="bold")
    ax1.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_:f"${v:,.0f}"))
    ax1.grid(axis="y",linestyle=":",alpha=0.35)
    ax2.bar(sls,ufs,color=cols,edgecolor="white")
    for i,v in enumerate(ufs):
        ax2.text(i,v+max(ufs)*0.015,f"${v:.2f}M",ha="center",va="bottom",fontsize=8)
    ax2.set_ylabel("Upfront cost, mid estimate ($M)",fontsize=9)
    ax2.set_title("XOS Hub — Nominal Upfront Cost by Site",fontsize=10,fontweight="bold")
    ax2.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_:f"${v:.1f}M"))
    ax2.grid(axis="y",linestyle=":",alpha=0.35)
    fig.tight_layout(); fig.savefig(path,dpi=150,bbox_inches="tight"); plt.close(fig)
    print(f"  {path.name}")

for sk,_ in SITES:
    fig_xos_daily(sk, FIGS/f"xos_{sk}_daily.png")
    fig_xos_breakdown(sk, FIGS/f"xos_{sk}_breakdown.png")
for sk in ["northgate","fresno"]:
    fig_kmp_daily(sk, FIGS/f"kmp_{sk}_daily.png")
    fig_kmp_breakdown(sk, FIGS/f"kmp_{sk}_breakdown.png")
fig_xsite_xos(FIGS/"xos_xsite_summary.png")
print()


# ─────────────────────────────────────────────────────────────────────
# WORD DOCUMENT HELPERS
# ─────────────────────────────────────────────────────────────────────
def _hdr_row(row):
    for cell in row.cells:
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
        shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"2F5496"); tcPr.append(shd)
        for p in cell.paragraphs:
            run=p.runs[0] if p.runs else p.add_run()
            run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); run.font.bold=True; run.font.size=Pt(9)

def _alt_row(row,even=False):
    if not even: return
    for cell in row.cells:
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
        shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"EEF2F8"); tcPr.append(shd)

def _set_col_width(table,col_idx,width_inches):
    for row in table.rows:
        cell=row.cells[col_idx]; tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        tcW=OxmlElement("w:tcW"); tcW.set(qn("w:w"),str(int(width_inches*1440))); tcW.set(qn("w:type"),"dxa")
        tcPr.append(tcW)

def add_table(doc,headers,rows,col_widths=None,font_size=9):
    t=doc.add_table(rows=1+len(rows),cols=len(headers)); t.style="Table Grid"
    hrow=t.rows[0]
    for i,h in enumerate(headers):
        hrow.cells[i].text=h
        for p in hrow.cells[i].paragraphs:
            for r in p.runs: r.font.size=Pt(font_size)
    _hdr_row(hrow)
    for ri,rdata in enumerate(rows):
        for ci,val in enumerate(rdata):
            cell=t.rows[ri+1].cells[ci]; cell.text=str(val)
            for p in cell.paragraphs:
                for r in p.runs: r.font.size=Pt(font_size)
        _alt_row(t.rows[ri+1],ri%2==0)
    if col_widths:
        for ci,w in enumerate(col_widths): _set_col_width(t,ci,w)
    doc.add_paragraph()
    return t

def add_fig(doc,path,caption,width=6.0):
    if not Path(path).exists():
        doc.add_paragraph(f"[Figure not found: {Path(path).name}]"); return
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path),width=Inches(width))
    cap=doc.add_paragraph(caption); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs: r.font.size=Pt(9); r.font.italic=True
    doc.add_paragraph()

def h1(doc,text):
    p=doc.add_heading(text,level=1); p.runs[0].font.size=Pt(13); return p
def h2(doc,text): return doc.add_heading(text,level=2)
def h3(doc,text): return doc.add_heading(text,level=3)
def body(doc,text,size=10):
    p=doc.add_paragraph(text)
    for r in p.runs: r.font.size=Pt(size)
    return p
def divider(doc,text):
    p=doc.add_paragraph()
    r=p.add_run(f"{'─'*10}  {text}  {'─'*10}")
    r.font.size=Pt(11); r.font.bold=True; r.font.color.rgb=RGBColor(0x2F,0x54,0x96)
    doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────
# BUILD DOCUMENT
# ─────────────────────────────────────────────────────────────────────
print("Building Word document …")
doc=Document()
sec=doc.sections[0]
sec.left_margin=sec.right_margin=Inches(1.0)
sec.top_margin=sec.bottom_margin=Inches(1.0)

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("DRAFT — FOR SHIMA'S REVIEW"); r.font.size=Pt(11); r.font.bold=True; r.font.color.rgb=RGBColor(0xC0,0,0)
doc.add_paragraph()
p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r2=p2.add_run("Appendix A: Cost Model Development"); r2.font.size=Pt(18); r2.font.bold=True
doc.add_paragraph()
p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
p3.add_run("Task 4482 — Mobile DC Fast Charging (MDCFC) Cost Analysis\n"
           "Caltrans Maintenance Stations | Report Date: June 30, 2026\n"
           "Analysis Period: May 2025 – April 2026").font.size=Pt(10)
doc.add_page_break()

# ── A.1 Overview ──────────────────────────────────────────────────────
h1(doc,"A.1  Overview and Scope")
body(doc,
    "This appendix presents the cost model developed for Task 4482 to evaluate two "
    "DC fast charging deployment options at four Caltrans maintenance stations. "
    "Results are organized in two parts: Part I covers the XOS Hub MC02 mobile DCFC; "
    "Part II covers the Kempower fixed DCFC "
    "(optimal charger mix selected by MILP). All costs use a 10-year asset lifespan "
    "and include capital amortization, TOU energy charges, utility demand/subscription "
    "charges, and warranty and maintenance (the latter two flagged as assumed values "
    "where the vendor contract price is not yet confirmed).")

add_table(doc,
    ["Site","State","Utility","Geotab Data Days","Kempower Status"],
    [["Northgate","Sacramento","SMUD C&I 21–299 kW","307","Complete"],
     ["Fresno","Fresno","PG&E BEV-2 Secondary","313","Complete — corrected PG&E rates"],
     ["Glendale","Glendale","PG&E BEV-2 (GWP proxy ⚠)","255","Complete — PG&E BEV-2 proxy"],
     ["San Diego","San Diego","SDG&E EV-HP Secondary","339","Paused — not yet computed"]],
    col_widths=[1.0,1.1,1.6,1.2,2.1])

# ── A.2 Data & Methodology ────────────────────────────────────────────
h1(doc,"A.2  Data Sources and Shared Methodology")
h2(doc,"A.2.1  Charging Event Data")
body(doc,"Charging events derived from Geotab zone-to-zone (Z2Z) trip records per site: "
     "vehicle ID, station arrival time, departure time, and estimated energy needed. "
     "Filtered to operating days May 2025 – April 2026. Dwell windows shorter than "
     "60 minutes are extended to 60 minutes.")

h2(doc,"A.2.2  TOU Energy Rates")
add_table(doc,
    ["Site","Utility","Season","Period","Hours (Pacific)","Rate ($/kWh)"],
    [["Northgate","SMUD","Summer (Jun–Sep)","Peak","16:00–21:00 weekdays","$0.2341"],
     ["Northgate","SMUD","Summer","Off-Peak","All other summer","$0.1215"],
     ["Northgate","SMUD","Non-Summer","Peak","16:00–21:00 weekdays","$0.1477"],
     ["Northgate","SMUD","Non-Summer","Saver","09:00–16:00","$0.0888"],
     ["Northgate","SMUD","Non-Summer","Off-Peak","All other","$0.1264"],
     ["Fresno","PG&E BEV-2","All","Peak","16:00–21:00, all days","$0.36977"],
     ["Fresno","PG&E BEV-2","All","Super Off-Peak","09:00–14:00, all days","$0.13327"],
     ["Fresno","PG&E BEV-2","All","Off-Peak","All other","$0.15654"],
     ["Glendale","PG&E BEV-2 (proxy A ⚠)","All","Peak","16:00–21:00, all days","$0.36977"],
     ["Glendale","PG&E BEV-2 (proxy A ⚠)","All","Super Off-Peak","09:00–14:00, all days","$0.13327"],
     ["Glendale","PG&E BEV-2 (proxy A ⚠)","All","Off-Peak","All other","$0.15654"],
     ["Glendale","SMUD C&I (proxy B ⚠)","Summer (Jun–Sep)","Peak","16:00–21:00 weekdays","$0.2341"],
     ["Glendale","SMUD C&I (proxy B ⚠)","Summer","Off-Peak","All other summer","$0.1215"],
     ["Glendale","SMUD C&I (proxy B ⚠)","Non-Summer","Peak","16:00–21:00 weekdays","$0.1477"],
     ["Glendale","SMUD C&I (proxy B ⚠)","Non-Summer","Saver","09:00–16:00","$0.0888"],
     ["Glendale","SMUD C&I (proxy B ⚠)","Non-Summer","Off-Peak","All other","$0.1264"],
     ["San Diego","SDG&E EV-HP","Summer (Jun–Oct)","On-Peak","16:00–21:00, all days","$0.29036"],
     ["San Diego","SDG&E EV-HP","Summer","Off-Peak","06:00–16:00","$0.12828"],
     ["San Diego","SDG&E EV-HP","Summer","Super Off-Peak","00:00–06:00","$0.12089"],
     ["San Diego","SDG&E EV-HP","Winter (Nov–May)","On-Peak","16:00–21:00, all days","$0.30199"],
     ["San Diego","SDG&E EV-HP","Winter","Off-Peak","06:00–16:00","$0.13067"],
     ["San Diego","SDG&E EV-HP","Winter","Super Off-Peak","00:00–06:00","$0.11588"]],
    col_widths=[0.85,1.2,1.1,1.1,1.5,1.05])

h2(doc,"A.2.3  Demand and Capacity Charges")
add_table(doc,
    ["Site","Utility","Charge Type","Rate ($/kW-mo)","Applied To","Notes"],
    [["Northgate","SMUD","Global demand charge","$6.454","All-hours peak grid draw","Both tiers apply"],
     ["Northgate","SMUD","Peak-window demand","$9.960","Peak-win peak kW (16:00–21:00 wkdays)",""],
     ["Fresno","PG&E BEV-2","Subscription charge","$1.91","Subscribed kW = peak draw","Overage: $3.82/kW"],
     ["Glendale","PG&E BEV-2 (proxy A ⚠)","Subscription charge","$1.91","Subscribed kW = peak draw","No peak-win demand tier"],
     ["Glendale","SMUD C&I (proxy B ⚠)","Global demand charge","$6.454","All-hours peak grid draw","Two-tier structure"],
     ["Glendale","SMUD C&I (proxy B ⚠)","Peak-window demand","$9.960","Peak-win peak kW (16:00–21:00 wkdays)","Two-tier structure"],
     ["San Diego","SDG&E EV-HP","Subscription charge","$4.81","Subscribed kW = peak draw","—"]],
    col_widths=[0.85,1.2,1.3,1.0,1.8,1.1])
body(doc,"All demand/subscription charges are amortized to a daily equivalent by dividing the "
     "monthly charge by 30.42 days/month and are INCLUDED in every cost total reported in this appendix.")
body(doc,"⚠ Glendale — two proxy rates analyzed (GWP actual tariff not confirmed): "
     "Proxy A (PG&E BEV-2) uses a subscription-only structure ($1.91/kW-mo, no peak-window tier) "
     "and higher TOU energy rates; Proxy B (SMUD C&I) uses a two-tier demand structure "
     "($6.454 + $9.960/kW-mo) and lower TOU energy rates. Both proxies are run through the full "
     "XOS simulation and Kempower MILP; results are compared in §A.13.")

h2(doc,"A.2.4  90th-Percentile Design Method")
body(doc,"The nominal design configuration for each site is determined by the 90th-percentile "
     "total daily cost method: (1) compute the all-in total daily cost for every operating day "
     "(capital + energy + demand + warranty + maintenance); (2) sort all days; (3) the value at "
     "the 90th percentile is the NOMINAL DESIGN COST; (4) the configuration (K or charger mix) "
     "on the day nearest that value is the NOMINAL DESIGN. This provides a cost-conservative "
     "target that covers all but the top 10% of days without over-provisioning for rare peaks.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PART I — XOS HUB MC02
# ══════════════════════════════════════════════════════════════════════
divider(doc,"PART I — XOS HUB MC02 (Mobile DC Fast Charging)")

# ── A.3 XOS Specs ─────────────────────────────────────────────────────
h1(doc,"A.3  XOS Hub MC02 — Hardware Specifications and Cost Assumptions")
add_table(doc,
    ["Parameter","Value","Source"],
    [["Model","XOS Hub MC02","XOS Trucks"],
     ["Form factor","Trailer-mounted, mobile DCFC","XOS User Manual"],
     ["Battery chemistry","Lithium Iron Phosphate (LFP)","XOS Manual Sec. 5"],
     ["Nominal battery capacity","282 kWh","XOS Manual Sec. 5"],
     ["Usable capacity (SoC ≥ 20%)","225.6 kWh","Derived: 80% of nominal"],
     ["Min. SoC reserve","20% (56.4 kWh)","Operational constraint"],
     ["Charge heads","4 × CCS1","XOS User Manual"],
     ["Max output per port","80 kW (constant)","XOS User Manual"],
     ["Max hub output — battery only","150 kW continuous","XOS User Manual"],
     ["Max hub output — grid + battery","230 kW continuous","XOS User Manual"],
     ["Max grid input per unit","~83 kW (480V × 100A × √3)","XOS User Manual"],
     ["Grid circuit requirement","480V, 3-phase, 100A dedicated","Electrical design"],
     ["Charge efficiency (grid→batt)","95%","Modeled assumption"],
     ["Discharge efficiency (batt→veh)","95%","Modeled assumption"],
     ["Service life","10 years / 3,000 cycles at 70% DoD","XOS Manual Sec. 5"],
     ["Time-step resolution","15 minutes","Model parameter"],
     ["Max units modeled per site","20","Model parameter"],
     ["Unit purchase price","$245,437.50","XOS Program Advisor (list price; not DGS contract)"],
     ["Annual maintenance *","$6,000 / unit / year","⚠ ASSUMED — no published data"],
     ["Annual warranty *","$10,000 / unit / year","Caltrans project team"]],
    col_widths=[2.2,2.0,2.8])
body(doc,"* Assumed values — must be confirmed with XOS before report finalization (see §A.14).")

# ── A.4 XOS Sizing Methodology ─────────────────────────────────────────
h1(doc,"A.4  XOS Hub MC02 — Sizing Methodology")
h2(doc,"A.4.1  Operating Logic")
body(doc,"Each XOS Hub unit remains connected to the 480V, 3-phase grid throughout the operating day. "
     "Grid power continuously replenishes the battery at up to 83 kW per unit. The battery acts as "
     "an energy buffer: vehicles draw from the battery while the grid recharges it simultaneously. "
     "SoC is held ≥ 20% at all times; if it drops below 20%, the Hub prioritizes grid recharge "
     "(proactive recharge mode) until SoC recovers. No diesel, no V2G, no disconnect events.")

h2(doc,"A.4.2  Adaptive-K Simulation")
body(doc,"For each operating day, the model iterates K from 1 upward, adding one XOS unit at a time "
     "and simulating the full day on 15-minute steps, until all arriving vehicles are fully served "
     "(energy need met within dwell window). The minimum K that achieves full service is recorded. "
     "Each unit dispatches vehicles to its 4 CCS1 charge heads on a greedy (first-come, "
     "first-served) basis at up to 80 kW per port.")

h2(doc,"A.4.3  Cost Computation")
add_table(doc,
    ["Cost Component","Formula","Daily $/unit at avg K"],
    [["Capital — purchase (amort.)","(Purchase × K) / (10 yr × 12 mo) / 30.42","Varies with K"],
     ["Capital — infra (amort.)","Infra_mid(K) / (10 yr × 12 mo) / 30.42","Varies with K"],
     ["Maintenance * (assumed)","($6,000 × K) / 12 / 30.42","$16.44 per unit"],
     ["Warranty * (assumed)","($10,000 × K) / 12 / 30.42","$27.41 per unit"],
     ["Energy","Σ grid_kW[t] × 0.25 h × TOU_rate[t]","Varies"],
     ["Demand — global","Peak_kW × rate / 30.42","Varies"],
     ["Demand — peak window","Peak_win_kW × rate / 30.42","SMUD/GWP only"]],
    col_widths=[1.8,2.8,1.8])

# ── A.5 XOS Per-Site Results ───────────────────────────────────────────
h1(doc,"A.5  XOS Hub MC02 — Per-Site Results")
fig_n=1
for sk,sl in SITES:
    d=XD[sk]; pday=d["pday"]; nomK=d["nomK"]; up=d["upfront"]
    h2(doc,f"A.5.{SITE_IDX[sk]+1}  {sl} ({UTILITY_SHORT[sk]})")
    if d["cap_hit"]:
        wp=doc.add_paragraph()
        wp.add_run("⚠ MODEL CAP REACHED — K=20 at p90 day. True requirement may exceed 20 units. "
                   "Costs are lower-bound estimates. See §A.14.").font.color.rgb=RGBColor(0xC0,0,0)

    # Key metrics table
    annual=d["dc"]["total_allin"].sum()*365/d["days"]
    add_table(doc,
        ["Metric","Value"],
        [["Utility",UTILITY_FULL[sk]],
         ["Analysis days",str(d["days"])],
         ["Avg. vehicles / day",f"{d['n_veh']:.1f}"],
         ["Avg. XOS units (K)",f"{d['sm']['K'].mean():.1f}"],
         ["K range (min–max)",f"{d['sm']['K'].min()} – {d['sm']['K'].max()}"],
         ["Most common K (mode)",str(d["sm"]["K"].mode()[0])],
         ["Avg. service rate",f"{d['svc_pct']:.1f}%"],
         ["90th-pct design cost",f"{fmt(d['p90'],2)}/day"],
         ["Nominal K at p90",f"{nomK} XOS Hub units"],
         ["Nominal design date",pday["date"].strftime("%Y-%m-%d")],
         ["Peak grid draw (nominal day)",f"{d['peak_kw']:.0f} kW"],
         ["Est. annual cost",f"{fmt(annual)}/yr"],
         ["Upfront — hardware",f"{fmt(up['hw'])}  (K × $245,438)"],
         ["Upfront — infra (mid)",f"{fmt(up['inf'])}  (building-side electrical)"],
         ["Upfront — TOTAL (mid)",fmt(up["total"])]],
        col_widths=[2.5,3.9])

    # Cost breakdown table
    r=pday
    gd_d=float(r["demand_global_monthly_$"])/DAYS_PER_MO
    pw_d=float(r["demand_peak_win_monthly_$"])/DAYS_PER_MO
    total_bd=float(r["purchase_capex_daily"])+float(r["infra_capex_daily"])+float(r["maint_daily"])+float(r["warranty_daily"])+float(r["energy_cost_daily"])+gd_d+pw_d
    add_table(doc,
        ["Cost Component","$/day — Nominal Design Day"],
        [["Capital — purchase (amort.)",fmt(float(r["purchase_capex_daily"]),2)],
         ["Capital — infrastructure (amort.)",fmt(float(r["infra_capex_daily"]),2)],
         ["Maintenance * (assumed)",fmt(float(r["maint_daily"]),2)],
         ["Warranty * (assumed)",fmt(float(r["warranty_daily"]),2)],
         ["Energy (TOU rate)",fmt(float(r["energy_cost_daily"]),2)],
         ["Demand — global (monthly ÷ 30.42)",fmt(gd_d,2)],
         ["Demand — peak window (monthly ÷ 30.42)",fmt(pw_d,2)],
         ["TOTAL (all-in)",fmt(total_bd,2)]],
        col_widths=[3.2,2.5])

    add_fig(doc,str(FIGS/f"xos_{sk}_daily.png"),
            f"Figure A.{fig_n}. {sl} — XOS Hub MC02 Daily Cost Distribution with 90th-Percentile Line.",width=6.2)
    fig_n+=1
    add_fig(doc,str(FIGS/f"xos_{sk}_breakdown.png"),
            f"Figure A.{fig_n}. {sl} — XOS Cost Breakdown at Nominal Design Day (K = {nomK}).",width=5.5)
    fig_n+=1
    if SITE_IDX[sk]<3: doc.add_page_break()

# ── A.6 XOS Cross-Site Summary ─────────────────────────────────────────
h1(doc,"A.6  XOS Hub MC02 — Cross-Site Summary and Configuration Ranking")
add_table(doc,
    ["Rank","Site","Nom. K","p90 Daily Cost","Est. Annual","Peak kW","Upfront (mid)","Flag"],
    sorted([
        [str(i+1),XD[sk]["label"],str(XD[sk]["nomK"]),
         f"{fmt(XD[sk]['p90'],2)}/day",
         fmt(XD[sk]["dc"]["total_allin"].sum()*365/XD[sk]["days"])+"/yr",
         f"{XD[sk]['peak_kw']:.0f} kW",
         fmt(XD[sk]["upfront"]["total"]),
         "⚠ K_CAP" if XD[sk]["cap_hit"] else ("⚠ GWP proxy" if sk=="glendale" else "—")]
        for i,(sk,_) in enumerate(sorted(SITES,key=lambda x:XD[x[0]]["upfront"]["total"]))
    ],key=lambda r:r[0]),
    col_widths=[0.4,0.9,0.55,1.1,1.1,0.75,1.1,0.95])

add_fig(doc,str(FIGS/"xos_xsite_summary.png"),
        f"Figure A.{fig_n}. XOS Hub MC02 — Cross-Site Comparison of p90 Design Cost and Upfront Cost.",width=6.2)
fig_n+=1

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PART II — KEMPOWER
# ══════════════════════════════════════════════════════════════════════
divider(doc,"PART II — KEMPOWER FIXED DCFC (MILP Optimal Charger Mix)")

# ── A.7 Kempower Specs & Pricing ───────────────────────────────────────
h1(doc,"A.7  Kempower Fixed DCFC — Hardware Specifications and Pricing")
body(doc,"Kempower charger pricing was received from the Caltrans ZEVI Statewide Program Advisor team "
     "under California DGS Contract 1-23-61-15A (National Car Charging LLC). These are state "
     "contract prices and serve as the authoritative cost basis.")

add_table(doc,
    ["Parameter","50 kW","150 kW","250 kW"],
    [["DGS Contract Group","Group 5","Group 6","Group 7"],
     ["DGS Model(s)","B-500 / B-501","S-600 / S-601","S-700 / S-701"],
     ["Power Output","50 kW DC","150 kW DC","250 kW DC"],
     ["Connector","CCS / CHAdeMO","CCS / CHAdeMO","CCS / CHAdeMO"],
     ["Form Factor","Portable cabinet","Fixed cabinet","Fixed cabinet"],
     ["Purchase Price (DGS)","$23,408 ¹","$62,154 ²","$101,946"],
     ["Installation (DGS)","$855","$4,750","$5,225"],
     ["Annual Maintenance (DGS)","$1,573/yr","$1,573/yr","$1,573/yr"],
     ["Annual Warranty *","$2,000/yr","$2,000/yr","$2,000/yr"],
     ["Service Life","8 years","8 years","8 years"],
     ["Daily CapEx (amort.)",fmt(KMP_DAILY["50kW"],2)+"/unit",
      fmt(KMP_DAILY["150kW"],2)+"/unit",fmt(KMP_DAILY["250kW"],2)+"/unit"]],
    col_widths=[2.0,1.6,1.6,1.6])
body(doc,"¹ Average of B-500 ($21,995) and B-501 ($24,820). "
     "² Average of S-600 ($62,113) and S-601 ($62,194). "
     "* Warranty assumed — to be confirmed. "
     "DGS contract allows 1%/2% discount for orders >$100K/>$500K (prices shown are pre-discount).")

# ── A.8 Kempower MILP Methodology ─────────────────────────────────────
h1(doc,"A.8  Kempower Fixed DCFC — MILP Sizing Methodology")
body(doc,"The Kempower charger mix is selected by a Mixed-Integer Linear Program (MILP) run "
     "independently for each operating day. The MILP simultaneously determines the number and "
     "type of Kempower chargers to deploy and the power schedule for each charging event "
     "across all 15-minute time intervals.")
add_table(doc,
    ["MILP Element","Description"],
    [["Decision variables","Count of 50 kW / 150 kW / 250 kW chargers (integers); power allocation to each vehicle at each time step"],
     ["Objective function","Minimize: capital amortization + TOU energy cost + demand/subscription charges"],
     ["Energy constraint","Each vehicle must receive required energy within its dwell window"],
     ["Power constraint","Total dispatch ≤ installed charger capacity at every time step"],
     ["Site-specific rates","Energy rates and demand coefficients set to site utility tariff — SMUD (Northgate), PG&E BEV-2 (Fresno + Glendale proxy), SDG&E EV-HP (San Diego)"],
     ["Solver","Gurobi (with HiGHS fallback)"],
     ["90th-pct design mix","Sort all days by all-in daily cost; the charger mix on the p90-closest day is the nominal design"]],
    col_widths=[1.8,4.5])

# ── A.9 Kempower Per-Site Results ──────────────────────────────────────
h1(doc,"A.9  Kempower Fixed DCFC — Per-Site Results")
kmp_sites_order = [("northgate","Northgate"),("fresno","Fresno"),("glendale","Glendale"),("san_diego","San Diego")]
for ksi,(sk,sl) in enumerate(kmp_sites_order,1):
    h2(doc,f"A.9.{ksi}  {sl}  —  Status: {KMP_SITE_STATUS[sk]}")
    kd=KD[sk]
    if kd is None:
        body(doc,f"Kempower results for {sl} are not yet available. "
             f"Status: {KMP_SITE_STATUS[sk]}. Results will be added when the pipeline completes.")
        continue

    summ=kd["summ"]; pday=kd["pday"]

    # Overview metrics table
    avg_demand_global = summ["global_demand_cost"].mean()/DAYS_PER_MO
    avg_demand_pw     = summ["peak_window_demand_cost"].mean()/DAYS_PER_MO
    avg_capex         = summ["capex_daily"].mean()
    avg_energy        = summ["energy_cost"].mean()
    avg_allin         = summ["total_allin"].mean()
    annual_kmp        = avg_allin*summ["date"].diff().dt.days.fillna(1).mean()*len(summ)
    annual_kmp2       = summ["total_allin"].sum()*365/len(summ)

    add_table(doc,
        ["Metric","Value"],
        [["Utility",UTILITY_FULL[sk]],
         ["Analysis days",str(len(summ))],
         ["Avg. vehicles / day",f"{summ['n_vehicles'].mean():.1f}"],
         ["Avg. chargers selected / day",f"{summ['n_chargers'].mean():.1f}"],
         ["Charger count range",f"{summ['n_chargers'].min()} – {summ['n_chargers'].max()}"],
         ["Avg. service rate",f"{summ['svc_rate_pct'].mean():.1f}%"],
         ["Avg. peak grid draw",f"{summ['peak_kw'].mean():.0f} kW"],
         ["90th-pct design cost (all-in)",f"{fmt(kd['p90'],2)}/day"],
         ["Nominal charger mix",str(pday["mix"])],
         ["Nominal design date",pday["date"].strftime("%Y-%m-%d")],
         ["Avg. daily CapEx",fmt(avg_capex,2)+"/day"],
         ["Avg. daily energy cost",fmt(avg_energy,2)+"/day"],
         ["Avg. demand — global (daily amort.)",fmt(avg_demand_global,2)+"/day"],
         ["Avg. demand — peak window (daily amort.)",fmt(avg_demand_pw,2)+"/day"],
         ["Avg. total all-in daily cost",fmt(avg_allin,2)+"/day"],
         ["Est. annual cost",fmt(annual_kmp2)+"/yr"]],
        col_widths=[2.5,3.9])

    # Cost breakdown at p90 day
    p_capex = float(pday["capex_daily"]); p_energy = float(pday["energy_cost"])
    p_gd = float(pday.get("global_demand_cost",0))/DAYS_PER_MO
    p_pw = float(pday.get("peak_window_demand_cost",0))/DAYS_PER_MO
    p_tot = p_capex+p_energy+p_gd+p_pw
    add_table(doc,
        ["Cost Component","$/day — Nominal Design Day"],
        [["Capital CapEx (amort.)",fmt(p_capex,2)],
         ["Energy (TOU rate)",fmt(p_energy,2)],
         ["Demand — global (monthly ÷ 30.42)",fmt(p_gd,2)],
         ["Demand — peak window (monthly ÷ 30.42)",fmt(p_pw,2)],
         ["TOTAL (all-in)",fmt(p_tot,2)]],
        col_widths=[3.2,2.5])

    # Most common mixes
    top_mixes=summ["mix"].value_counts().head(6)
    add_table(doc,
        ["Charger Mix","Days","% of days"],
        [[mix,str(cnt),f"{100*cnt/len(summ):.1f}%"] for mix,cnt in top_mixes.items()],
        col_widths=[3.5,0.7,0.9])

    # Add figures for complete sites
    if sk in ("northgate","fresno"):
        add_fig(doc,str(FIGS/f"kmp_{sk}_daily.png"),
                f"Figure A.{fig_n}. {sl} — Kempower DCFC Daily Cost Distribution with 90th-Percentile Line.",width=6.2)
        fig_n+=1
        add_fig(doc,str(FIGS/f"kmp_{sk}_breakdown.png"),
                f"Figure A.{fig_n}. {sl} — Kempower Cost Breakdown at Nominal Design Day.",width=5.5)
        fig_n+=1

    doc.add_page_break()

# ── A.10 Kempower Partial Summary ──────────────────────────────────────
h1(doc,"A.10  Kempower Fixed DCFC — Partial Cross-Site Summary")
body(doc,"Table A.10.1 summarizes available Kempower results. San Diego results are paused. "
     "Glendale results use PG&E BEV-2 proxy (Proxy A); see §A.13 for SMUD proxy comparison.")
kmp_summary_rows=[]
for sk,sl in SITES:
    kd=KD[sk]
    if kd is None:
        kmp_summary_rows.append([sl,UTILITY_SHORT[sk],"—","—","—","—","—",KMP_SITE_STATUS[sk]])
    else:
        avg_allin=kd["summ"]["total_allin"].mean()
        annual=avg_allin*365
        kmp_summary_rows.append([
            sl,UTILITY_SHORT[sk],
            f"{kd['summ']['n_chargers'].mean():.1f}",
            f"{fmt(kd['p90'],2)}/day",
            fmt(annual)+"/yr",
            f"{kd['summ']['peak_kw'].mean():.0f} kW",
            str(kd["pday"]["mix"]),
            KMP_SITE_STATUS[sk]])
add_table(doc,
    ["Site","Utility","Avg K","p90 Daily Cost","Est. Annual","Peak kW","p90 Mix","Status"],
    kmp_summary_rows,
    col_widths=[0.85,1.1,0.6,1.1,1.0,0.7,1.5,1.3])

# ── A.11 Northgate XOS vs Kempower ────────────────────────────────────
doc.add_page_break()
h1(doc,"A.11  Northgate — XOS Hub vs. Kempower Comparison")
body(doc,"Table A.11.1 provides a head-to-head comparison of both scenarios for Northgate, "
     "the only site where complete results exist for both technologies.")
xos_n=XD["northgate"]; kmp_n=KD["northgate"]
kmp_annual=kmp_n["summ"]["total_allin"].mean()*365 if kmp_n else None
xos_annual=xos_n["dc"]["total_allin"].sum()*365/xos_n["days"]
add_table(doc,
    ["Metric","XOS Hub MC02","Kempower Fixed DCFC"],
    [["Utility","SMUD C&I","SMUD C&I"],
     ["Nominal design K / mix",f"{xos_n['nomK']} XOS units",str(kmp_n["pday"]["mix"]) if kmp_n else "—"],
     ["90th-pct daily cost",fmt(xos_n["p90"],2)+"/day",fmt(kmp_n["p90"],2)+"/day" if kmp_n else "—"],
     ["Est. annual cost",fmt(xos_annual)+"/yr",fmt(kmp_annual)+"/yr" if kmp_annual else "—"],
     ["Avg. peak grid draw",f"{xos_n['dc']['peak_grid_kw'].mean():.0f} kW",f"{kmp_n['summ']['peak_kw'].mean():.0f} kW" if kmp_n else "—"],
     ["Upfront cost (hardware only)",fmt(xos_n["upfront"]["hw"]),"Varies daily (per-day MILP)"],
     ["Upfront — infra (mid)",fmt(xos_n["upfront"]["inf"]),"Included in purchase/install"],
     ["Upfront — TOTAL (mid)",fmt(xos_n["upfront"]["total"]),"Varies by mix"],
     ["Asset lifespan","10 years","8 years"],
     ["Mobility","Mobile — can be relocated","Fixed — permanent installation"],
     ["Grid infrastructure","480V, 3-phase, 100A / unit","Direct grid connection / charger"],
     ["Avg. service rate",f"{xos_n['svc_pct']:.1f}%",f"{kmp_n['summ']['svc_rate_pct'].mean():.1f}%" if kmp_n else "—"]],
    col_widths=[2.2,2.2,2.2])

# ── A.12 Findings ─────────────────────────────────────────────────────
h1(doc,"A.12  Findings and Recommendations")
findings={
    "Northgate (XOS)":
        f"Nominal design: K = {XD['northgate']['nomK']} XOS Hub units at a 90th-pct daily cost of "
        f"{fmt(XD['northgate']['p90'],2)}/day. SMUD demand charges are significant: the two-tier "
        f"SMUD structure adds ~${XD['northgate']['pday']['demand_peak_win_monthly_$']/DAYS_PER_MO:.0f}/day "
        f"(peak-window) on top of the global demand charge. The high peak grid draw "
        f"({XD['northgate']['peak_kw']:.0f} kW) should be confirmed against the available electrical "
        f"service capacity at the site.",
    "Northgate (Kempower)":
        f"The MILP-selected optimal mix varies by day. The most common configurations are 5-charger "
        f"mixes combining 50 kW, 150 kW, and 250 kW units. The 90th-pct daily cost of "
        f"{fmt(kmp_n['p90'],2)}/day" if kmp_n else "—" + " is substantially lower than XOS at Northgate, "
        "primarily because the Kempower units require no battery buffer and their amortized capital "
        "costs are lower on a per-kW basis." if kmp_n else
        f"Kempower Northgate results are in progress.",
    "Fresno (XOS)":
        f"Nominal design: K = {XD['fresno']['nomK']} units at {fmt(XD['fresno']['p90'],2)}/day. "
        f"PG&E's subscription model (no peak-window demand) simplifies the cost structure. "
        f"Fresno has the lowest average fleet size ({XD['fresno']['n_veh']:.1f} veh/day).",
    "Glendale":
        f"Nominal design: K = {XD['glendale']['nomK']} XOS units at {fmt(XD['glendale']['p90'],2)}/day "
        f"(PG&E BEV-2 proxy A). Two utility proxies were analyzed since GWP's actual tariff is not "
        f"confirmed — see §A.13 for the side-by-side comparison. "
        f"The SMUD proxy (higher demand charges) yields higher costs; the PG&E BEV-2 proxy "
        f"(subscription-only demand, higher TOU energy) yields lower costs. "
        f"Glendale has the lightest fleet of the four sites.",
    "San Diego (XOS)":
        f"Highest-demand site ({XD['san_diego']['n_veh']:.1f} veh/day average). K=20 model cap "
        f"is binding at the p90 day — true design requirement likely exceeds 20 units. "
        f"Cost figures are lower-bound estimates only. An alternative sizing strategy is recommended.",
}
for title,text in findings.items():
    p=doc.add_paragraph()
    r=p.add_run(f"{title}: "); r.bold=True; r.font.size=Pt(10)
    r2=p.add_run(str(text)); r2.font.size=Pt(10)

# ── A.13 Glendale Utility Rate Sensitivity ────────────────────────────
doc.add_page_break()
h1(doc,"A.13  Glendale — Utility Rate Sensitivity (SMUD vs. PG&E BEV-2 Proxy)")
body(doc,
    "Because Glendale Water & Power's actual commercial tariff (GWP Schedule LD-2/PC-1) could not "
    "be obtained, two California utility proxies were analyzed to bracket the likely cost range. "
    "Proxy A uses PG&E BEV-2 Secondary (subscription-based demand, higher peak TOU energy). "
    "Proxy B uses SMUD C&I 21–299 kW (two-tier demand charges, lower TOU energy). "
    "XOS simulation outputs (K, grid draw profiles) are identical under both proxies since the "
    "simulation is not rate-aware; only the cost calculation differs. "
    "Kempower MILP results may differ between proxies because the rate structure enters the "
    "optimization objective directly.")

# Build summary rows
def _gl_xos_row(label, d, smud_note=""):
    if d is None:
        return [label,"—","—","—","—","—","Pending"]
    pday=d["pday"]
    gd_daily=float(pday["demand_global_monthly_$"])/DAYS_PER_MO
    pw_daily=float(pday["demand_peak_win_monthly_$"])/DAYS_PER_MO
    return [label, str(d["nomK"]),
            fmt(d["p90"],2)+"/day",
            fmt(float(pday["energy_cost_daily"]),2),
            fmt(gd_daily,2),
            fmt(pw_daily,2),
            fmt(d["p90"],2)+"/day"]

def _gl_kmp_row(label, kd):
    if kd is None:
        return [label,"—","—","—","—","—","—","Pending"]
    pday=kd["pday"]
    gd=float(pday.get("global_demand_cost",0))/DAYS_PER_MO
    pw=float(pday.get("peak_window_demand_cost",0))/DAYS_PER_MO
    return [label, str(pday["mix"]),
            fmt(kd["p90"],2)+"/day",
            fmt(float(pday["capex_daily"]),2),
            fmt(float(pday["energy_cost"]),2),
            fmt(gd,2),
            fmt(pw,2),
            fmt(kd["p90"],2)+"/day"]

h2(doc,"A.13.1  XOS Hub MC02 — Rate Sensitivity")
add_table(doc,
    ["Rate Proxy","Nominal K","p90 Daily Cost","Energy Cost","Demand — Global","Demand — Peak Win","All-In Total"],
    [_gl_xos_row("PG&E BEV-2 (Proxy A)", XD["glendale"]),
     _gl_xos_row("SMUD C&I (Proxy B)", XD_GL_SMUD)],
    col_widths=[1.5,0.8,1.2,1.1,1.3,1.3,1.2])
body(doc,"Note: Nominal K is identical under both proxies (XOS simulation is rate-independent). "
     "Cost differences are driven entirely by rate structure: SMUD's two-tier demand "
     "charges dominate at high peak draws while PG&E's subscription charge is lower but "
     "peak TOU energy rate is higher ($0.370/kWh vs. $0.234/kWh).")

h2(doc,"A.13.2  Kempower Fixed DCFC — Rate Sensitivity")
add_table(doc,
    ["Rate Proxy","Nominal Mix","p90 Daily Cost","CapEx","Energy","Demand — Global","Demand — Peak Win","All-In Total"],
    [_gl_kmp_row("PG&E BEV-2 (Proxy A)", KD["glendale"]),
     _gl_kmp_row("SMUD C&I (Proxy B)", KD_GL_SMUD)],
    col_widths=[1.4,1.6,1.1,0.8,0.8,1.1,1.1,1.0])
body(doc,"Note: Kempower nominal mix may differ between proxies because demand charges enter the "
     "MILP objective directly, influencing the optimal charger configuration. "
     "SMUD's peak-window demand charge ($9.960/kW-mo) incentivizes spreading load across more "
     "chargers to reduce peak kW; PG&E's subscription charge ($1.91/kW-mo) has a weaker effect. "
     "Glendale has the lightest vehicle demand of all four sites (1–3 chargers on most days), "
     "so the mix difference is modest.")

h2(doc,"A.13.3  Cross-Technology Summary (Glendale)")
gl_xos_pge = XD["glendale"]; gl_xos_smud = XD_GL_SMUD
gl_kmp_pge = KD["glendale"]; gl_kmp_smud = KD_GL_SMUD
rows_cross = []
for tech, xd_p, xd_s, kd_p, kd_s in [
    ("XOS — p90 daily cost",
     fmt(gl_xos_pge["p90"],2)+"/day" if gl_xos_pge else "—",
     fmt(gl_xos_smud["p90"],2)+"/day" if gl_xos_smud else "—",
     "n/a","n/a"),
    ("XOS — nominal K",
     str(gl_xos_pge["nomK"]) if gl_xos_pge else "—",
     str(gl_xos_smud["nomK"]) if gl_xos_smud else "—",
     "n/a","n/a"),
    ("Kempower — p90 daily cost",
     "n/a","n/a",
     fmt(gl_kmp_pge["p90"],2)+"/day" if gl_kmp_pge else "Pending",
     fmt(gl_kmp_smud["p90"],2)+"/day" if gl_kmp_smud else "Pending"),
    ("Kempower — nominal mix",
     "n/a","n/a",
     str(gl_kmp_pge["pday"]["mix"]) if gl_kmp_pge else "Pending",
     str(gl_kmp_smud["pday"]["mix"]) if gl_kmp_smud else "Pending"),
]:
    rows_cross.append([tech, xd_p, xd_s, kd_p, kd_s])
add_table(doc,
    ["Metric","XOS — PG&E (A)","XOS — SMUD (B)","KMP — PG&E (A)","KMP — SMUD (B)"],
    rows_cross,
    col_widths=[2.2,1.3,1.3,1.3,1.3])
body(doc,"⚠ All Glendale figures are estimates. Obtain GWP Schedule LD-2/PC-1 tariff "
     "to replace these proxies. Contact GWP Customer Service: 855-550-4497.")

# ── A.14 Open Items ────────────────────────────────────────────────────
h1(doc,"A.14  Open Items and Assumptions Requiring Confirmation")
open_items=[
    ("XOS maintenance pricing","$6,000/unit/yr assumed — no published service contract data. Confirm with XOS Program Advisor."),
    ("XOS warranty pricing","$10,000/unit/yr provided by Caltrans team — confirm current pricing and applicability."),
    ("XOS unit price — DGS status","$245,437.50 is a list price from XOS Program Advisor, not a CA DGS contract. Clarify whether a DGS contract price is available."),
    ("Glendale — GWP utility tariff","GWP Schedule LD-2/PC-1 tariff was not obtainable. PG&E BEV-2 is used as a proxy (effective 2026-07-01). Replace with actual GWP tariff when available — contact GWP Customer Service: 855-550-4497."),
    ("San Diego — K cap binding","True design requirement at p90 day likely exceeds K=20 model ceiling. Re-run with raised cap or evaluate hybrid strategy."),
    ("SDG&E super off-peak hours","Assumed 00:00–06:00. Confirm exact super off-peak hours with SDG&E."),
    ("End-of-life / decommissioning","Not included in current cost model. Obtain estimates from XOS and Kempower vendors."),
    ("Kempower warranty","$2,000/yr per charger assumed — confirm with Kempower or Caltrans contract."),
    ("Glendale Kempower — PG&E proxy rerun","Kempower MILP rerun with PG&E BEV-2 proxy rates completed 2026-07-01. All Glendale figures reflect PG&E rates. Replace with actual GWP tariff when available."),
    ("San Diego Kempower","Not yet computed. Pipeline is paused pending direction on site priority."),
]
add_table(doc,
    ["Item","Description"],
    open_items,
    col_widths=[2.0,4.4])

# ── A.15 Worst-Day Tables ──────────────────────────────────────────────
h1(doc,"A.15  Worst-Day Supporting Tables (XOS Hub MC02)")
add_table(doc,
    ["Site","Date","K","Total Daily Cost","Energy","Demand-Global","Demand-PkWin","Peak kW"],
    [[sl,
      XD[sk]["dc"].loc[XD[sk]["dc"]["total_allin"].idxmax(),"date"].strftime("%Y-%m-%d"),
      str(int(XD[sk]["dc"].loc[XD[sk]["dc"]["total_allin"].idxmax(),"K"])),
      fmt(float(XD[sk]["dc"]["total_allin"].max()),2),
      fmt(float(XD[sk]["dc"].loc[XD[sk]["dc"]["total_allin"].idxmax(),"energy_cost_daily"]),2),
      fmt(float(XD[sk]["dc"].loc[XD[sk]["dc"]["total_allin"].idxmax(),"demand_global_monthly_$"])/DAYS_PER_MO,2),
      fmt(float(XD[sk]["dc"].loc[XD[sk]["dc"]["total_allin"].idxmax(),"demand_peak_win_monthly_$"])/DAYS_PER_MO,2),
      f"{float(XD[sk]['dc'].loc[XD[sk]['dc']['total_allin'].idxmax(),'peak_grid_kw']):.0f} kW"
     ] for sk,sl in SITES],
    col_widths=[0.9,1.0,0.4,1.1,0.9,1.1,1.1,0.85])

doc.save(str(DOCOUT))
print(f"\nSaved: {DOCOUT}")
print(f"Figures: {FIGS}")
